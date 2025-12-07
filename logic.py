import docx
import fitz  # PyMuPDF
import re
import zipfile
import io
import os
import logging
import concurrent.futures
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image, ImageEnhance  # Tambah ImageEnhance untuk preprocessing
from typing import List, Dict, Any, Union, Pattern, Tuple, Callable, Optional

# --- KONFIGURASI LOGGING ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

# --- KONFIGURASI OCR ---
# Tetap batasi halaman agar tidak timeout, tapi kualitas per halaman dimaksimalkan
MAX_OCR_PAGES = 50 

# --- 1. DEFINISI POLA REGEX ---

BASE_PATTERNS: Dict[str, List[str]] = {
    "Fintech": [r"fintech", r"financial technology"],
    "AI": [r"ai"],
    "Artificial Intelligence": [r"artificial intelligence"],
    "Blockchain": [r"blockchain"],
    "Big Data": [r"big data"],
    "Machine Learning": [r"machine learning"],
    "Cloud Computing": [r"cloud computing"],
    "Digital Banking": [r"digital banking"],
    "Mobile Banking": [r"mobile banking"],
    "E-Payment": [r"e-payment", r"electronic payment"],
    "Cybersecurity": [r"cybersecurity", r"cyber security"],
    "Digital Service": [r"digital service"]
}

INDO_PATTERNS: Dict[str, List[str]] = {
    "Fintech": [r"teknologi finansial"],
    "AI": [r"ai"],
    "Artificial Intelligence": [r"kecerdasan buatan"],
    "Blockchain": [r"rantai blok"],
    "Big Data": [r"mahadata"],
    "Machine Learning": [r"pembelajaran mesin"],
    "Cloud Computing": [r"komputasi awan"],
    "Digital Banking": [r"layanan perbankan digital", r"bank digital"],
    "Mobile Banking": [r"m-banking"],
    "E-Payment": [r"pembayaran elektronik", r"pembayaran digital"],
    "Cybersecurity": [r"keamanan siber"],
    "Digital Service": [r"layanan digital"]
}

def get_compiled_patterns(bilingual: bool = False) -> Dict[str, Pattern]:
    compiled = {}
    for key, patterns in BASE_PATTERNS.items():
        combined_patterns = patterns.copy()
        if bilingual and key in INDO_PATTERNS:
            combined_patterns.extend(INDO_PATTERNS[key])
        regex_str = r"\b(" + "|".join(combined_patterns) + r")\b"
        compiled[key] = re.compile(regex_str, re.IGNORECASE)
    return compiled

PATTERNS_EN = get_compiled_patterns(bilingual=False)
PATTERNS_BI = get_compiled_patterns(bilingual=True)


# --- 2. FUNGSI PEMBACAAN FILE ---

def extract_year(filename: str, zip_name: str) -> str:
    year_pattern = r'\b(20\d{2}|19\d{2})\b'
    match_file = re.search(year_pattern, filename)
    if match_file: return match_file.group(1)
    match_zip = re.search(year_pattern, zip_name)
    if match_zip: return match_zip.group(1)
    return "Unknown"

def read_txt(file_bytes: bytes) -> str:
    try: return file_bytes.decode("utf-8")
    except UnicodeDecodeError: return file_bytes.decode("latin-1")

def read_docx(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        tables = [cell.text for t in doc.tables for r in t.rows for cell in r.cells]
        return '\n'.join(paragraphs + tables)
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        return ""

def preprocess_image_for_ocr(img: Image.Image) -> Image.Image:
    """
    Meningkatkan kualitas gambar untuk hasil OCR yang lebih baik pada dokumen scan/buram.
    """
    # 1. Konversi ke Grayscale (L)
    img = img.convert('L')
    
    # 2. Tingkatkan Kontras (Membantu memisahkan teks pudar dari background)
    enhancer_contrast = ImageEnhance.Contrast(img)
    img = enhancer_contrast.enhance(2.0)  # Naikkan kontras 2x lipat
    
    # 3. Tingkatkan Ketajaman (Membantu scan yang blur)
    enhancer_sharpness = ImageEnhance.Sharpness(img)
    img = enhancer_sharpness.enhance(2.0) # Naikkan ketajaman 2x lipat
    
    return img

def read_pdf_ocr(file_bytes: bytes, filename: str) -> str:
    """
    OCR Robust: Menggunakan DPI tinggi dan Preprocessing Citra.
    """
    logger.info(f"[{filename}] Memulai proses OCR Enhanced...")
    text_content = []
    try:
        # PENTING: Naikkan DPI ke 300 untuk menangkap huruf kecil.
        # Ini akan sedikit lebih lambat dari 200, tapi jauh lebih akurat.
        images = convert_from_bytes(
            file_bytes, 
            dpi=300, 
            fmt='jpeg', 
            grayscale=True, # Awal grayscale dari poppler
            last_page=MAX_OCR_PAGES
        )
        
        total_pages_scanned = len(images)
        logger.info(f"[{filename}] Mengonversi {total_pages_scanned} halaman pertama (300 DPI).")
        
        for i, img in enumerate(images):
            if (i + 1) % 5 == 0:
                logger.info(f"[{filename}] OCR processing page {i+1}/{total_pages_scanned}")
            
            # --- TAHAP PREPROCESSING CITRA ---
            # Proses gambar agar lebih tajam dan kontras tinggi sebelum dibaca Tesseract
            processed_img = preprocess_image_for_ocr(img)
            
            # Gunakan config standard psm 3 (Auto)
            text = pytesseract.image_to_string(processed_img, lang='eng+ind', config='--psm 3') 
            text_content.append(text)
            
        logger.info(f"[{filename}] OCR selesai ({total_pages_scanned} halaman).")
        return '\n'.join(text_content)
    except Exception as e:
        logger.error(f"[{filename}] OCR Gagal: {e}")
        return ""

def read_pdf(file_bytes: bytes, filename: str, include_scanned: bool = False) -> Tuple[str, bool]:
    """
    Returns: Tuple[str, bool] -> (Extracted Text, Is_Skipped)
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        full_text = ""
        for page in doc:
            full_text += page.get_text()
        
        # Logika Deteksi Scan
        char_count = len(full_text.strip())
        page_count = len(doc)
        avg_chars = char_count / page_count if page_count > 0 else 0
        
        is_scanned = avg_chars < 50 
        
        if is_scanned:
            logger.warning(f"[{filename}] Terdeteksi sebagai SCAN (Avg chars: {avg_chars:.1f}).")
            
            if not include_scanned:
                logger.info(f"[{filename}] SKIPPING file karena opsi 'Include Scanned' tidak dicentang.")
                return "", True 
            
            return read_pdf_ocr(file_bytes, filename), False
            
        logger.info(f"[{filename}] Ekstraksi teks PDF native berhasil.")
        return full_text, False
        
    except Exception as e:
        logger.warning(f"[{filename}] Gagal baca native PDF: {e}. Mencoba OCR jika diizinkan.")
        if include_scanned:
            return read_pdf_ocr(file_bytes, filename), False
        return "", True

def count_total_words(text: str) -> int:
    if not text: return 0
    return len(text.split())


# --- 3. LOGIKA UTAMA (ANALISIS) ---

def analyze_single_file(args: Tuple) -> Dict[str, Any]:
    zip_name, filename, file_bytes, is_bilingual, include_scanned = args
    
    logger.info(f"[{filename}] Memulai analisis...")
    
    patterns = PATTERNS_BI if is_bilingual else PATTERNS_EN
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    is_skipped = False
    
    if ext == '.pdf': 
        text, is_skipped = read_pdf(file_bytes, filename, include_scanned)
    elif ext == '.docx': 
        text = read_docx(file_bytes)
    elif ext == '.txt': 
        text = read_txt(file_bytes)
    
    bank_name = os.path.splitext(zip_name)[0]
    year = extract_year(filename, zip_name)
    
    if is_skipped:
        return {
            "Nama Bank": bank_name,
            "Tahun": year,
            "Nama File": os.path.basename(filename),
            "Total Kata Dokumen": 0,
            "Status": "SKIPPED (Scan)"
        }

    total_words = count_total_words(text)
    
    counts = {}
    for key, pattern in patterns.items():
        matches = pattern.findall(text)
        counts[key] = len(matches)
    
    logger.info(f"[{filename}] Analisis selesai. Total kata: {total_words}")
    
    row_data = {
        "Nama Bank": bank_name,
        "Tahun": year,
        "Nama File": os.path.basename(filename),
        "Total Kata Dokumen": total_words,
        "Status": "SUCCESS"
    }
    row_data.update(counts)
    
    return row_data

def process_zip_file(
    uploaded_zip: Any, 
    is_bilingual: bool = False, 
    include_scanned: bool = False, 
    progress_callback: Optional[Callable[[int, int, str], None]] = None
) -> List[Dict[str, Any]]:
    
    logger.info(f"Membuka ZIP: {uploaded_zip.name}")
    files_to_process = []
    
    with zipfile.ZipFile(uploaded_zip) as z:
        all_files = z.namelist()
        valid_extensions = ['.pdf', '.docx', '.txt']
        
        target_files = [
            f for f in all_files 
            if not f.endswith('/') and '__MACOSX' not in f and os.path.splitext(f)[1].lower() in valid_extensions
        ]
        
        logger.info(f"Ditemukan {len(target_files)} file valid untuk diproses.")
        
        for filename in target_files:
            with z.open(filename) as f:
                content = f.read()
                files_to_process.append((uploaded_zip.name, filename, content, is_bilingual, include_scanned))
    
    results_list = []
    total_files = len(files_to_process)
    
    if total_files == 0:
        return []

    with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
        future_to_file = {
            executor.submit(analyze_single_file, args): args[1] 
            for args in files_to_process
        }
        
        completed_count = 0
        for future in concurrent.futures.as_completed(future_to_file):
            filename = future_to_file[future]
            try:
                data = future.result()
                results_list.append(data)
            except Exception as e:
                logger.error(f"[{filename}] Exception unhandled: {e}")
                results_list.append({
                    "Nama File": os.path.basename(filename),
                    "Status": "ERROR",
                    "Total Kata Dokumen": 0
                })
            
            completed_count += 1
            if progress_callback:
                progress_callback(completed_count, total_files, filename)
                
    logger.info("Semua file selesai diproses.")
    return results_list

def generate_csv_output(results_list: List[Dict[str, Any]]) -> str:
    if not results_list: return ""
    keyword_headers = list(BASE_PATTERNS.keys())
    header = ["Nama Bank", "Tahun", "Nama File", "Status"] + keyword_headers + ["Total Kata Dokumen"]
    csv_str = ",".join(header) + "\n"
    for row in results_list:
        line = []
        line.append(f'"{row.get("Nama Bank", "")}"')
        line.append(f'"{row.get("Tahun", "")}"')
        line.append(f'"{row.get("Nama File", "")}"')
        line.append(f'"{row.get("Status", "UNKNOWN")}"')
        
        if row.get("Status") == "SKIPPED (Scan)":
            for _ in keyword_headers: line.append("0")
        else:
            for kw in keyword_headers:
                line.append(str(row.get(kw, 0)))
                
        line.append(str(row.get("Total Kata Dokumen", 0)))
        csv_str += ",".join(line) + "\n"
    return csv_str